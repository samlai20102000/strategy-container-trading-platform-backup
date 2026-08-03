from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
EVIDENCE = ROOT / "artifacts" / "backtest-100dd"
ASSETS = EVIDENCE / "report-assets"
REPORT_DIR = EVIDENCE / "final-report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

JOB_SUMMARY_PATH = EVIDENCE / "job_1785770356467_b7fe7008_forensic_summary.json"
EXCEL_SUMMARY_PATH = EVIDENCE / "testing2_forensic_summary.json"
ASSESSMENT_PATH = EVIDENCE / "persisted_job_risk_integrity_assessment.json"
CHART_EVIDENCE_PATH = ASSETS / "chart_evidence.json"
MATRIX_PATH = EVIDENCE / "nine_strategy_validation_matrix.md"

JOB_SUMMARY = json.loads(JOB_SUMMARY_PATH.read_text(encoding="utf-8"))
EXCEL_SUMMARY = json.loads(EXCEL_SUMMARY_PATH.read_text(encoding="utf-8"))
ASSESSMENT = json.loads(ASSESSMENT_PATH.read_text(encoding="utf-8"))
CHART_EVIDENCE = json.loads(CHART_EVIDENCE_PATH.read_text(encoding="utf-8"))

JOB = JOB_SUMMARY["job"]
METRICS = JOB["metrics"]
ACCOUNTING = JOB["accounting"]
PATH_FORENSICS = JOB_SUMMARY["pathForensics"]
TRADE_CALC = EXCEL_SUMMARY["independent_trade_recalculation"]
EXPOSURE = EXCEL_SUMMARY["exposure_forensics"]
RISK = ASSESSMENT["assessment"]

DOCX_PATH = REPORT_DIR / "回測中心_100%最大回撤_最終法證分析與全策略優化報告.docx"
MD_PATH = REPORT_DIR / "回測中心_100%最大回撤_最終法證分析與全策略優化報告.md"


def money(value: float | int) -> str:
    return f"{value:,.2f} USDT"


def pct(value: float | int, digits: int = 2) -> str:
    return f"{value:,.{digits}f}%"


def multiple(value: float | int, digits: int = 2) -> str:
    return f"{value:,.{digits}f} 倍"


def iso_date(ms: int | float) -> str:
    return datetime.fromtimestamp(float(ms) / 1000, tz=timezone.utc).strftime(
        "%Y-%m-%d %H:%M UTC"
    )


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell: Any, inches: float) -> None:
    width = int(inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))
    cell.width = Inches(inches)


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell: Any, top: int = 85, start: int = 100, bottom: int = 85, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table: Any, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for index, width in enumerate(widths):
            set_cell_width(row.cells[index], width)
            set_cell_margins(row.cells[index])
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float]) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, "1F4E78")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(value))
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "EAF2F8")
        for column_index, value in enumerate(values):
            paragraph = cells[column_index].paragraphs[0]
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8.5)

    set_fixed_table_layout(table, widths)
    doc.add_paragraph()
    return table


def add_paragraph(doc: Document, text: str = "", bold_prefix: str | None = None) -> Any:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet(doc: Document, text: str, level: int = 0) -> Any:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_number(doc: Document, text: str) -> Any:
    paragraph = doc.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_callout(doc: Document, title: str, body: str, fill: str = "FFF2CC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_fixed_table_layout(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(128, 0, 0)
    paragraph.add_run(f"\n{body}")
    doc.add_paragraph()


def add_toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "開啟 Word 後更新欄位以顯示目錄"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, placeholder, end):
        run._r.append(node)


def add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    paragraph.add_run(" 頁")


def set_picture_alt_text(inline_shape: Any, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc: Document, path: Path, caption: str, alt: str) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(path), width=Inches(6.4))
    set_picture_alt_text(shape, caption, alt)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_paragraph.paragraph_format.space_after = Pt(8)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    for name, size, color in (
        ("Title", 24, "17365D"),
        ("Subtitle", 12, "666666"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "1F4E78"),
        ("Heading 3", 11, "244062"),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.space_before = Pt(10 if name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        styles[list_style].font.name = "Arial"
        styles[list_style]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("回測中心 100% 最大回撤\n最終法證分析與全策略優化報告")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("KAMA_RAINBOW_MARTIN_V1｜BTC-USDT｜30m｜SINGLE_EXCLUSIVE")

    doc.add_paragraph()
    add_table(
        doc,
        ["欄位", "內容"],
        [
            ("事故 job", JOB["jobId"]),
            ("回測 run", "bt_KAMARAINBOWMARTINV1_1785770363974_tx4"),
            ("資料期間", f"{JOB['startDateIso'][:10]} 至 {JOB['endDateIso'][:10]}"),
            ("報告日期", "2026-08-04"),
            ("證據狀態", "原始交易檔、持久化 job、equityCurve、程式碼與測試均已交叉驗證"),
        ],
        [1.55, 4.95],
    )
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice.add_run("機密技術報告｜不得把本報告中的失效回測 KPI 用作實盤決策依據")
    notice_run.bold = True
    notice_run.font.color.rgb = RGBColor(192, 0, 0)
    doc.add_page_break()


def add_exec_summary(doc: Document) -> None:
    doc.add_heading("執行摘要", level=1)
    add_callout(
        doc,
        "最終判定",
        "畫面顯示的 100% 最大回撤不是百分比公式或 UI 欄位錯綁；在原 runner 產生的逐 K mark-to-market 權益路徑中，帳戶確實穿越 0。真正的系統錯誤，是該 S1 runner 宣告已套用曝險、保證金與清算政策，實際卻未在訂單准入與逐 K 估值中執行，並在破產後繼續持倉，最終把不可生存路徑包裝為 completed 高獲利回測。",
        "FCE4D6",
    )

    add_table(
        doc,
        ["問題", "法證結論", "處置"],
        [
            (
                "100% 最大回撤是否算錯",
                "否。權益曾由正峰值跌至負值；有限責任口徑下回撤封頂 100% 正確。",
                "保留公式，不得把 100% 人工改小。",
            ),
            (
                "結果是否仍可視為有效策略績效",
                "否。破產後恢復、未宣告 bankrupt、零次清算及政策超限同時存在。",
                "結果必須 fail closed，不顯示為可信 KPI。",
            ),
            (
                "testing2.xlsx 為何看不到 100% 回撤",
                "該檔只有 91 筆平倉交易，沒有逐 K 未實現損益；已實現權益最大回撤僅 1.82 USDT。",
                "交易表與 equityCurve 必須分開標示與匯出。",
            ),
            (
                "上一輪費用修正是否正確",
                "否。open position 應扣已發生的入場費，而不是尚未發生的出場費；已恢復 canonical 口徑。",
                "保留專門回歸測試，禁止再次改壞。",
            ),
            (
                "本次是否只修 KRM",
                "否。發布邊界守門與 risk-aware 降採樣接在所有九個內建 runner 共用 finalizer／三個採樣入口。",
                "九策略 source contract 與完整測試套件已通過。",
            ),
        ],
        [1.45, 3.35, 1.70],
    )

    add_paragraph(
        doc,
        "最重要的產品語義是：本次優化不會把一條已破產路徑「美化」成較小回撤。當同類路徑再次發生時，系統應以專用風險完整性錯誤拒絕發布，而不是回傳 100% 回撤但仍標成成功。若要取得真正可交易的修正後 KPI，下一階段必須把 legacy S1 runner 遷移到逐單准入與逐 K 清算的 runtime risk kernel 後重新回測。",
    )


def add_scope_and_method(doc: Document) -> None:
    doc.add_heading("一、分析範圍與證據方法", level=1)
    add_paragraph(
        doc,
        "本報告採用四層對帳：交易檔、持久化回測 job、逐 K 權益路徑與程式執行鏈。任何結論至少由兩個獨立來源支持；無法從平倉交易表推導的風險事件，不以推測填補。",
    )

    doc.add_heading("1.1 主要證據", level=2)
    add_table(
        doc,
        ["證據", "用途", "完整性"],
        [
            ("testing2.xlsx", "91 筆平倉交易、PnL、數量、馬丁層與名義本金", f"SHA-256 {EXCEL_SUMMARY['source']['sha256'][:16]}…"),
            ("持久化 job JSON", "原始 metrics、accounting、executionPolicy、資料品質", JOB["jobId"]),
            ("持久化 equityCurve", "查明首次非正權益、最低權益、恢復路徑", "2,001 點舊版報告曲線"),
            ("原始 runner／finalizer 程式碼", "確認風控是執行邏輯或只有 metadata", "逐函式稽核"),
            ("獨立重算與真實資料重播", "排除 UI、Excel 與單一公式誤差", "不寫入交易／不觸發實盤 mutation"),
            ("Vitest、TypeScript、production build", "驗證修正未破壞既有契約", "最終完整回歸"),
        ],
        [1.55, 3.10, 1.85],
    )

    doc.add_heading("1.2 口徑分離", level=2)
    add_bullet(doc, "已實現權益：只在平倉時累加交易淨損益；testing2.xlsx 能重建此口徑。")
    add_bullet(doc, "逐 K mark-to-market 權益：已實現權益加上持倉毛浮盈虧及已發生費用；最大回撤與夏普使用此口徑。")
    add_bullet(doc, "有限責任權益：帳戶資產不得低於 0；到達清算／破產條件後必須終止或維持 0，禁止恢復。")
    add_bullet(doc, "策略局部止損：價格或策略狀態條件，不等於交易所保證金與強制清算。")


def add_numeric_reconciliation(doc: Document) -> None:
    doc.add_heading("二、數值法證與完整對帳", level=1)
    doc.add_heading("2.1 畫面、Excel 與 accounting 對帳", level=2)
    add_table(
        doc,
        ["項目", "獨立重算／原始值", "對帳說明"],
        [
            ("平倉交易數", f"{TRADE_CALC['trade_count']} 筆", "與畫面 90 勝／1 負一致"),
            ("平倉 PnL 合計", money(TRADE_CALC["pnl_sum_usdt"]), "Excel 91 筆逐列加總"),
            ("期末未實現 PnL", money(ACCOUNTING["unrealizedPnl"]), "尚有 1 個 long open position"),
            ("總回報 USDT", money(METRICS["totalReturnUSDT"]), "20,861.23 − 2.15 = 20,859.08"),
            ("期末權益", money(ACCOUNTING["finalEquity"]), "10,000 + 20,859.08 = 30,859.08"),
            ("已實現權益最大回撤", f"{money(TRADE_CALC['max_drawdown_closed_trade_equity_usdt'])}／{pct(TRADE_CALC['max_drawdown_closed_trade_equity_pct'], 4)}", "只能反映平倉序列，不能代表持倉期間風險"),
            ("逐 K 最大回撤", f"{pct(METRICS['maxDrawdown'])}／{money(METRICS['maxDrawdownUSDT'])}", "來自完整 runner 權益路徑"),
        ],
        [1.65, 1.80, 3.05],
    )

    add_callout(
        doc,
        "為何 Excel 與畫面看似矛盾",
        "Excel 的 91 筆資料只記錄交易平倉結果。第 29 個 cycle 在持倉途中曾產生約 −13,802.57 USDT 的浮動毛損，令權益穿越 0，但價格後來回落，該 cycle 最終以 +7,241.75 USDT 平倉。因此平倉交易表呈現大幅獲利，卻無法否定期間內已發生破產。",
        "DDEBF7",
    )

    doc.add_heading("2.2 事故時間線", level=2)
    first_non_positive = PATH_FORENSICS["firstNonPositiveEquity"]
    minimum = PATH_FORENSICS["minimumEquity"]
    peak = PATH_FORENSICS["maximumDrawdownPeak"]
    add_table(
        doc,
        ["UTC 時間", "權益／事件", "法證意義"],
        [
            (peak["timestampIso"].replace("T", " ").replace(".000Z", " UTC"), money(peak["equity"]), "100% 回撤事件的前置正峰值（舊版降採樣曲線）"),
            ("2025-05-22 00:00 UTC", money(753.91), "尚未歸零，但只剩峰值的 6.61%"),
            (first_non_positive["timestampIso"].replace("T", " ").replace(".000Z", " UTC"), money(first_non_positive["equity"]), "首次非正權益；有限責任帳戶應在此前清算／破產"),
            (minimum["timestampIso"].replace("T", " ").replace(".000Z", " UTC"), money(minimum["equity"]), "舊版持久化曲線最低權益"),
            ("2025-05-23 18:00 UTC", money(3130.25), "破產後重新變正，直接證明 runner 繼續持有未清算部位"),
            ("2025-06-01 16:00 UTC", money(7241.75), "第 29 cycle 最終平倉獲利，掩蓋中途已破產"),
        ],
        [1.70, 1.80, 3.00],
    )

    add_figure(
        doc,
        ASSETS / "equity_insolvency_path.png",
        "圖 1　原始持久化權益曲線：正峰值、首次穿越零、最低權益與恢復",
        "BTC-USDT 30 分鐘 KRM 回測權益曲線，標示 2025 年 5 月的峰值、首次負權益、最低權益及破產後恢復。",
    )

    doc.add_heading("2.3 為何畫面是 11,450.29，但舊持久化曲線重算為 11,401.08", level=2)
    add_paragraph(
        doc,
        "畫面指標在 runner 的完整 27,745 點權益路徑上先計算，再把曲線等距降採樣至 2,001 點保存。舊等距抽樣恰好漏掉完整路徑的精確峰值，所以畫面最大回撤金額為 11,450.29 USDT，而從已保存曲線獨立重算為 11,401.08 USDT，差額 49.21 USDT。兩者百分比均為 100%，因為谷值已低於 0。這不是最大回撤公式錯誤，而是舊版曲線持久化的風險極值保真度不足；本次已以 risk-aware 降採樣修補。",
    )


def add_root_cause(doc: Document) -> None:
    doc.add_heading("三、唯一根因與失效機制", level=1)
    doc.add_heading("3.1 馬丁幾何級數使名義曝險遠超資本", level=2)
    add_paragraph(
        doc,
        "本次設定首單 100 USDT、倍率 2、最多 15 層。若把第 n 個訂單視為一個層級，累積名義本金為 100 × (2ⁿ − 1)。因此第 6 層為 6,300 USDT、第 7 層為 12,700 USDT，而 12 個訂單（交易紀錄 martinLayer=11）已達 409,500 USDT。",
    )
    add_table(
        doc,
        ["風險量", "宣告上限", "原事故觀察值", "超限程度"],
        [
            ("初始資金", money(JOB["initialCapital"]), money(JOB["initialCapital"]), "基準"),
            ("S1 gross notional", "權益 100%", money(RISK["observedEntryNotionalPeak"]), multiple(EXPOSURE["maximum_entry_notional_to_initial_capital"])),
            ("1× 槓桿 margin usage", "權益 40%", money(RISK["observedEntryMarginPeak"]), "相對初始 margin 上限約 102.38 倍"),
            ("保守全期峰值 gross 上限", money(RISK["grossNotionalLimitAtGlobalPeak"]), money(RISK["observedEntryNotionalPeak"]), multiple(RISK["observedEntryNotionalPeak"] / RISK["grossNotionalLimitAtGlobalPeak"])),
            ("保守全期峰值 margin 上限", money(RISK["marginLimitAtGlobalPeak"]), money(RISK["observedEntryMarginPeak"]), multiple(RISK["observedEntryMarginPeak"] / RISK["marginLimitAtGlobalPeak"])),
        ],
        [1.65, 1.45, 1.70, 1.70],
    )

    add_figure(
        doc,
        ASSETS / "martin_layer_risk_breach.png",
        "圖 2　100 USDT、2 倍馬丁的累積名義本金與政策門檻",
        "對數尺度顯示第 6 層首次超過 40% margin budget，第 7 層首次超過 100% gross-notional budget，最終達約 419,517 USDT。",
    )

    doc.add_heading("3.2 宣告的 policy 只有 metadata，沒有進入 KRM 訂單准入", level=2)
    add_paragraph(
        doc,
        "持久化 executionContext 宣告 riskModelVersion=gross-margin-liquidation-v2、leverage=1、maxGrossNotionalPct=100、maxMarginUsagePct=40；然而 KRM 專屬 runner 的模擬帳戶未接入訂單前風險判斷，新增 layer 時只依策略管理核心決策與 maxLayers 執行。結果是第 6／7 層本應被政策拒絕，卻一路成交至 409,500 USDT。",
    )

    doc.add_heading("3.3 hardStopLossPct=8% 不是清算保護", level=2)
    add_paragraph(
        doc,
        "在 40.95 倍名義曝險下，只需約 2.44% 的逆向價格變動就可耗盡 10,000 USDT 初始資本；8% 價格型 hard stop 理論虧損可達約 32,760 USDT，遠晚於破產。因此策略止損即使正常運作，也不能替代 max gross、margin admission、maintenance margin 與 liquidation。",
    )

    doc.add_heading("3.4 破產後未清算，造成不可能的績效組合", level=2)
    add_table(
        doc,
        ["違約代碼", "實際證據", "含義"],
        [
            (
                violation["code"],
                (money(violation["actual"]) if "actual" in violation else "路徑事件")
                + (f"；上限 {money(violation['limit'])}" if "limit" in violation else ""),
                violation["message"],
            )
            for violation in RISK["violations"]
        ],
        [2.05, 1.65, 2.80],
    )

    add_paragraph(
        doc,
        "這五類違約共同解釋畫面：總回報與 profit factor 使用最終平倉／期末帳本；最大回撤與夏普使用逐 K mark-to-market 曲線。由於 runner 允許負權益後恢復，兩組指標可以同時存在，但整份結果已失去有限責任與可執行性，不能被視為有效回測。",
    )


def add_formula_audit(doc: Document) -> None:
    doc.add_heading("四、最大回撤公式與上一輪錯誤修正稽核", level=1)
    doc.add_heading("4.1 最大回撤公式本身正確", level=2)
    add_paragraph(
        doc,
        "產品使用正峰值到谷值的 drawdown：DD(t) = [Peak(t) − max(Equity(t), 0)] ÷ Peak(t)。只要 Equity(t) ≤ 0 且歷史 Peak(t) > 0，有限責任回撤即為 100%。因此移除 100% 上限或改用絕對值都會掩蓋破產，不是修正。",
    )

    doc.add_heading("4.2 上一輪把 entry fee 改成 exit fee 是錯誤", level=2)
    add_paragraph(
        doc,
        "現有帳本在開倉時尚未把入場費從 realized equity 扣除，所以 open-position snapshot 必須以毛浮盈虧減去 entryFees，才能把已發生成本納入 mark-to-market。把它改成估計 exitFees 會同時漏掉已發生費用並提前計入未發生成本，且 canonical continuity test 已直接顯示 19.80 與 19.78 的回歸差異。本次已恢復 entry-fee 口徑。該 0.02 USDT 等級差異與 11,450.29 USDT 回撤沒有因果關係。",
    )

    add_callout(
        doc,
        "更正聲明",
        "上一輪『入場費重複扣除造成負權益』的根因判斷不成立。負權益來自超額名義曝險與缺少清算；費用口徑只是後續被引入的次生回歸，現已回復並受測試保護。",
        "FCE4D6",
    )


def add_implementation(doc: Document) -> None:
    doc.add_heading("五、本次已實施的全策略優化", level=1)
    add_table(
        doc,
        ["層級", "已實施內容", "行為變化"],
        [
            ("P0－會計恢復", "open-position unrealized PnL 恢復扣 entryFees 的 canonical 語義", "消除上一輪次生回歸"),
            ("P0－共用 fail-closed", "在所有 runner 共用 finalizer 內執行 backtest-risk-integrity-v1", "異常結果在 canonical artifact 保存前拋出專用錯誤"),
            ("P0－五類風險守門", "負權益、缺少 bankrupt、破產後恢復、gross 超限、margin 超限", "不再把失效路徑標為 completed KPI"),
            ("P0－risk-aware 降採樣", "保留首尾、全域高低、最大回撤峰谷、首次非正、前一點與首次恢復", "2,001 點曲線不再漏掉風險極值／瞬時破產"),
            ("P0－九策略 source contract", "九個內建 S1 runner 必須經過 finalizeV25Result 與 assertBacktestRiskIntegrity", "未來新增直接 return runner 結果會使測試失敗"),
            ("P0－結構化錯誤", "將風險完整性失敗納入回測錯誤分類", "前端／任務層可區分資料風險與一般執行錯誤"),
        ],
        [1.25, 3.30, 1.95],
    )

    add_figure(
        doc,
        ASSETS / "risk_integrity_guard_behavior.png",
        "圖 3　修正後發布邊界：異常回測 fail closed，不再產生可信 KPI",
        "流程圖比較舊版 completed KPI 路徑與新版風險完整性守門拒絕路徑。",
    )

    doc.add_heading("5.1 全部九個策略的覆蓋", level=2)
    strategy_rows = [
        ("strategy_20415", "runRainbow20415Backtest", "共用 finalizer"),
        ("RAINBOW_TREND_LADDER_V1", "runRainbowTrendLadderBacktest", "共用 finalizer＋risk-aware sampling"),
        ("KAMA_RAINBOW_MARTIN_V1", "runKamaRainbowMartinBacktest", "共用 finalizer＋risk-aware sampling"),
        ("KAMA_3K_BREAKOUT_V25", "runV25Backtest", "共用 finalizer"),
        ("20415_KAMA_MARTIN_V35", "共用 KAMA runner", "共用 finalizer＋risk-aware sampling"),
        ("20415_KAMA_MARTIN_V41", "共用 KAMA runner", "共用 finalizer＋risk-aware sampling"),
        ("KAMA_3K_ULTIMATE_V50", "共用 KAMA runner", "共用 finalizer＋risk-aware sampling"),
        ("KAMA_3K_HF_V61", "共用 KAMA runner", "共用 finalizer＋risk-aware sampling"),
        ("KAMA_3K_TORNADO_V70", "runV70Backtest", "共用 finalizer"),
    ]
    add_table(doc, ["策略 key", "S1 執行路徑", "本次覆蓋"], strategy_rows, [2.45, 2.30, 1.75])

    doc.add_heading("5.2 本次修正的誠實邊界", level=2)
    add_callout(
        doc,
        "目前是安全阻斷，不是完整經濟重算",
        "POSTHOC_ONLY 守門可防止錯誤 KPI 發布，但它不會替 legacy runner 回填『如果第 6／7 層當時被拒單』後的替代交易路徑。要得到修正後真正的收益、勝率與回撤，必須先把 runner 接入 runtime risk kernel，再以同一資料重跑。",
        "FFF2CC",
    )


def add_validation(doc: Document) -> None:
    doc.add_heading("六、驗證結果與驗收矩陣", level=1)
    add_table(
        doc,
        ["驗證項", "結果", "判定"],
        [
            ("原事故 job 套用新守門", "passed=false；5 類違約", "正確拒絕"),
            ("風險完整性單元測試", "4/4 通過", "規則語義正確"),
            ("risk-aware 降採樣", "4/4 通過", "瞬時破產與 DD 峰谷必保留"),
            ("九策略派發契約", "2/2 通過", "禁止繞過共用 finalizer"),
            ("聚焦風險回歸", "3 files／10 tests 通過", "守門、覆蓋、採樣共同正確"),
            ("完整 Vitest", "142 files passed；1,129 tests passed；2 files／5 tests skipped", "無已知測試回歸"),
            ("TypeScript", "tsc --noEmit 通過", "型別契約一致"),
            ("Production build", "Vite＋server bundle 通過", "可部署；僅既有 chunk-size 警告"),
        ],
        [2.00, 2.80, 1.70],
    )

    doc.add_heading("6.1 修正後應見行為", level=2)
    add_number(doc, "使用相同 KRM 設定重跑時，若 legacy runner 仍生成負權益／超額曝險，job 應失敗並回傳風險完整性錯誤，不應再顯示 completed KPI。")
    add_number(doc, "已存在的歷史 job 是法證證據，不會被本次程式修改回寫；在 UI 重新開啟舊 job 時仍可能看到原 100% 指標，這不代表新守門未生效。")
    add_number(doc, "完成 runtime kernel 遷移後，同一資料重跑必須在第 6 層 margin 或第 7 層 gross 門檻前拒單，且不得出現負權益或破產後恢復。")


def add_roadmap(doc: Document) -> None:
    doc.add_heading("七、完整優化路線與優先順序", level=1)
    add_table(
        doc,
        ["優先級", "工作", "驗收標準", "狀態"],
        [
            ("P0", "恢復 canonical fee、全策略 fail closed、risk-aware sampling", "原事故 job 被拒絕；全測試與 build 通過", "已完成"),
            ("P1", "把 KRM／Rainbow／legacy KAMA S1 遷移至 threeModePortfolioKernel 或同一 runtime admission API", "每次 order 前校驗 gross／margin；逐 K maintenance liquidation；bankruptcy terminal", "待實施"),
            ("P1", "為歷史 job 增加 read-time integrity badge／INVALID 狀態遷移", "舊失效 job 不再以一般 completed 卡片呈現；不刪除原始證據", "待實施"),
            ("P1", "完整 risk event ledger 與 Excel 風險工作表", "每層准入、拒單、margin、liq price、破產事件可追溯", "待實施"),
            ("P2", "把 hard stop 與清算距離並列顯示", "任何策略止損晚於 liquidation 時，preflight 明確拒絕或警告", "建議"),
            ("P2", "禁止無效 KPI 排名／比較", "riskIntegrity.passed=false 時不計入策略排名、快照或優化器目標", "建議"),
            ("P3", "大型前端 chunk 拆分", "降低目前約 3.52 MB JS bundle 與冷啟動成本", "非本事故阻斷項"),
        ],
        [0.65, 2.70, 2.45, 0.70],
    )

    doc.add_heading("7.1 Runtime kernel 的最低必要契約", level=2)
    add_bullet(doc, "Order admission：以成交後 gross notional、margin usage、maxOpenLegs 與方向互斥政策判斷，超限不得改寫策略持倉狀態。")
    add_bullet(doc, "Mark-to-market：每根已收盤 K 線計算 wallet balance、unrealized PnL、equity、used margin、maintenance margin 與 headroom。")
    add_bullet(doc, "Liquidation：equity ≤ maintenance margin 時強制關閉；資產以有限責任封底為 0；bankrupt=true 且狀態不可逆。")
    add_bullet(doc, "State synchronization：只有模擬交易所確認成交後，策略管理狀態才可增加 layer；拒單不得造成 ghost layer。")
    add_bullet(doc, "Evidence：所有拒單、清算與破產事件寫入 canonical ledger，績效與 Excel 從同一 ledger 派生。")

    doc.add_heading("7.2 完成 P1 後的數值驗收", level=2)
    add_table(
        doc,
        ["驗收項", "必須成立"],
        [
            ("曝險", "observed gross notional ≤ policy gross limit；used margin ≤ policy margin limit"),
            ("權益", "minimum equity ≥ 0；若 bankrupt，之後所有點維持 0"),
            ("狀態", "marginLiquidationCount 與 ledger 事件一致；禁止 POST_INSOLVENCY_RECOVERY"),
            ("資料", "完整 metrics 與持久化 risk-aware curve 的最大回撤峰谷完全一致"),
            ("UI", "INVALID／FAILED 不顯示為綠色成功卡，也不參與策略比較"),
            ("回歸", "九策略各至少一組正常、拒單、清算、破產 terminal、切片連續性測試"),
        ],
        [2.05, 4.45],
    )


def add_final_answer(doc: Document) -> None:
    doc.add_heading("八、最終回答", level=1)
    add_paragraph(
        doc,
        "最大回撤之所以仍是 100%，是因為原回測並沒有被錯誤公式誤判；它真的在 2025-05-22 把模擬權益打到負值。系統真正的 bug，是 KRM 的 SINGLE_EXCLUSIVE runner 沒有執行它宣告的 max gross、margin usage 與 liquidation policy，讓 100 USDT、2 倍馬丁一路擴張到約 409,500 USDT 名義倉位，並在破產後繼續持倉直至回復獲利。",
    )
    add_paragraph(
        doc,
        "因此正確處理不是把 100% 改小，而是把該結果判定為無效並禁止發布。本次已完成全九策略共用 fail-closed、五類風險完整性檢查、canonical fee 恢復與 risk-aware 權益降採樣；原事故 job 已被新守門判定不合格。下一個必要工程是 runtime risk kernel 遷移，完成後才能重新產生具有經濟意義的『修正後回撤與收益』。在此之前，原畫面的 +208.59%、98.9% 勝率、profit factor 與夏普均不得作為策略優劣或實盤配置依據。",
    )


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("附錄 A｜關鍵證據雜湊與檔案", level=1)
    evidence_paths = [
        Path("/home/ubuntu/upload/testing2.xlsx"),
        JOB_SUMMARY_PATH,
        EXCEL_SUMMARY_PATH,
        ASSESSMENT_PATH,
        EVIDENCE / "job_1785770356467_b7fe7008_equity_curve.json",
        EVIDENCE / "job_1785770356467_b7fe7008_trades.json",
        ROOT / "server/services/backtest/backtestRiskIntegrity.ts",
        ROOT / "server/services/backtest/equityCurveDownsample.ts",
        EVIDENCE / "final_vitest.log",
        EVIDENCE / "final_build.log",
    ]
    rows = []
    for path in evidence_paths:
        if path.exists():
            rows.append((path.name, sha256(path), f"{path.stat().st_size:,} bytes"))
    add_table(doc, ["檔案", "SHA-256", "大小"], rows, [2.25, 3.35, 0.90])

    doc.add_heading("附錄 B｜已修改與新增的核心檔案", level=1)
    add_table(
        doc,
        ["檔案", "用途"],
        [
            ("backtestRiskIntegrity.ts", "全策略 posthoc fail-closed 與五類違約診斷"),
            ("backtestEngine.ts", "共用 finalizer 接線、保存前守門、risk metadata"),
            ("backtestContracts.ts", "恢復 entry-fee 估值並保存 riskIntegrity 契約"),
            ("backtestRunnerPreflight.ts", "風險完整性結構化錯誤分類"),
            ("equityCurveDownsample.ts", "保留破產與最大回撤關鍵點的共用降採樣"),
            ("kamaRainbowMartinBacktest.ts", "改用共用 risk-aware sampler"),
            ("rainbowTrendLadderBacktest.ts", "改用共用 risk-aware sampler"),
            ("backtestRiskIntegrity*.test.ts", "規則、九策略覆蓋與 source contract"),
            ("equityCurveDownsample.test.ts", "瞬時破產／最大回撤峰谷／三入口接線測試"),
        ],
        [2.85, 3.65],
    )

    doc.add_heading("附錄 C｜限制與風險揭露", level=1)
    add_bullet(doc, "原完整 27,745 點 runner 曲線未直接持久化；原 job 保存的是 2,001 點舊等距抽樣曲線。完整 metrics 仍保留 11,450.29 USDT，而舊曲線可重建 11,401.08 USDT。")
    add_bullet(doc, "原事故中負權益點幸運地仍被舊抽樣保留；其他 job 可能漏掉瞬時破產，因此本次新增 risk-aware sampler。")
    add_bullet(doc, "本次未刪除或改寫歷史 job；歷史數據保留供稽核。")
    add_bullet(doc, "本次未觸發任何實盤下單、撤單、帳戶轉帳或交易所 mutation。")
    add_bullet(doc, "本報告是軟體與回測模型法證分析，不構成投資建議。")


def build_docx() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header.paragraphs[0]
    header.text = "回測中心 100% 最大回撤｜最終法證分析與全策略優化報告"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    doc.add_heading("目錄", level=1)
    add_toc(doc)
    doc.add_page_break()
    add_exec_summary(doc)
    add_scope_and_method(doc)
    add_numeric_reconciliation(doc)
    add_root_cause(doc)
    add_formula_audit(doc)
    add_implementation(doc)
    add_validation(doc)
    add_roadmap(doc)
    add_final_answer(doc)
    add_appendices(doc)

    core = doc.core_properties
    core.title = "回測中心 100% 最大回撤最終法證分析與全策略優化報告"
    core.subject = "KAMA_RAINBOW_MARTIN_V1 回測風險完整性"
    core.author = "Manus AI"
    core.keywords = "backtest, drawdown, risk integrity, martingale, liquidation"
    core.comments = "以 testing2.xlsx、持久化 job、equityCurve 與程式碼完成的法證報告"
    doc.save(DOCX_PATH)


def build_markdown() -> None:
    violations = "\n".join(
        f"| `{item['code']}` | {item['message']} | {item.get('actual', '路徑事件')} | {item.get('limit', '—')} |"
        for item in RISK["violations"]
    )
    markdown = f"""# 回測中心 100% 最大回撤：最終法證分析與全策略優化報告

> **最終判定：** 100% 最大回撤不是公式或 UI 顯示錯誤。原 runner 的逐 K mark-to-market 權益確實穿越 0；真正 bug 是 S1 KRM runner 沒有執行已宣告的曝險、保證金與清算政策，並在破產後繼續持倉與恢復獲利。

## 核心證據

| 項目 | 數值 |
|---|---:|
| 初始資金 | {money(JOB['initialCapital'])} |
| 總回報 | {pct(METRICS['totalReturn'])}／{money(METRICS['totalReturnUSDT'])} |
| 平倉交易 | {TRADE_CALC['trade_count']} 筆；{TRADE_CALC['winning_trades']} 勝／{TRADE_CALC['losing_trades']} 負 |
| 平倉 PnL 合計 | {money(TRADE_CALC['pnl_sum_usdt'])} |
| 期末未實現 PnL | {money(ACCOUNTING['unrealizedPnl'])} |
| 最大回撤 | {pct(METRICS['maxDrawdown'])}／{money(METRICS['maxDrawdownUSDT'])} |
| 首次非正權益 | {money(PATH_FORENSICS['firstNonPositiveEquity']['equity'])}，{PATH_FORENSICS['firstNonPositiveEquity']['timestampIso']} |
| 最低權益 | {money(PATH_FORENSICS['minimumEquity']['equity'])} |
| 最大入場名義本金 | {money(RISK['observedEntryNotionalPeak'])}（{multiple(EXPOSURE['maximum_entry_notional_to_initial_capital'])}初始資本） |

## Excel 與畫面的對帳

`testing2.xlsx` 只有 91 筆平倉交易，重算 PnL 為 {money(TRADE_CALC['pnl_sum_usdt'])}。期末尚有一個 open position，未實現 PnL 為 {money(ACCOUNTING['unrealizedPnl'])}，所以畫面總回報為 {money(METRICS['totalReturnUSDT'])}。Excel 的已實現權益最大回撤僅 {money(TRADE_CALC['max_drawdown_closed_trade_equity_usdt'])}，因為它沒有逐 K 持倉浮盈虧；100% 回撤來自另一份 mark-to-market equityCurve。

第 29 cycle 在持倉途中產生約 {money(abs(PATH_FORENSICS['activeClosedTradeRecordsAtEvent'][0]['markGrossPnlAtEvent']))} 浮動毛損，權益跌至負值，但價格後來回落，最終以 {money(PATH_FORENSICS['activeClosedTradeRecordsAtEvent'][0]['pnl'])} 平倉。這就是「90 勝 1 負、總回報 +208.59%，但最大回撤 100%」可以同時出現的原因；整條路徑在有限責任模型下無效。

## 五類風險完整性違約

| 代碼 | 訊息 | 實際 | 上限 |
|---|---|---:|---:|
{violations}

## 已完成修正

1. 恢復 open-position 扣 `entryFees` 的 canonical 估值，撤銷上一輪錯誤修改。
2. 在全部九個 runner 的共用 finalizer 加入 `backtest-risk-integrity-v1` fail-closed。
3. 結果在 canonical artifact 保存前檢查負權益、bankruptcy 狀態、破產後恢復、gross 與 margin 超限。
4. 新增 risk-aware equity downsampling，永遠保留最大回撤峰谷、最低權益、首次非正與首次恢復點。
5. 新增九策略 source contract 與專門測試；完整結果為 142 files、1,129 tests 通過，production build 通過。

## 仍需完成的 P1

目前修正是 **POSTHOC_ONLY 安全阻斷**，不會虛構一組較漂亮的修正後 KPI。下一階段必須把 legacy S1 runner 遷移到共用 runtime risk kernel，於每次加倉前執行 gross／margin admission，並於每根 K 線執行 maintenance-margin liquidation 與 terminal bankruptcy。完成後才可用同一資料重跑，取得真正具有經濟意義的新收益與回撤。

## 最終回答

最大回撤仍為 100% 的根本原因是**回測帳戶真的被超額馬丁曝險打穿**；不是回撤公式錯誤。正確修正是拒絕發布失效結果並把風控放進 runtime，而不是把 100% 人工調小。原畫面的收益、勝率、profit factor 與夏普不得作為實盤配置依據。
"""
    MD_PATH.write_text(markdown, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(json.dumps({
        "docx": str(DOCX_PATH),
        "docx_sha256": sha256(DOCX_PATH),
        "markdown": str(MD_PATH),
        "markdown_sha256": sha256(MD_PATH),
    }, ensure_ascii=False, indent=2))
